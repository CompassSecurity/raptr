"""
Unit tests for the report render layer.

Focus: the Jinja2 SandboxedEnvironment must block SSTI/RCE payloads on both the
HTML and DOCX paths while still rendering benign templates correctly.
"""

import io
from datetime import datetime

import pytest
from docx import Document
from jinja2.exceptions import SecurityError

from app.services.report.render import render_docx_report, render_html_report
from app.services.report.report_data import AssessmentInfo, ReportContext

# Classic Jinja2 SSTI -> RCE payload (harmless `id` command).
RCE_PAYLOAD = (
    "{{ self.__init__.__globals__.__builtins__.__import__('os').popen('id').read() }}"
)


def _context() -> ReportContext:
    """Minimal, DB-free ReportContext for rendering."""
    return ReportContext(
        assessment=AssessmentInfo(
            id="1", name="Acme Engagement", description="", assessment_type="RedTeam"
        ),
        activities_grouped=[],
        activities_flat=[],
        statistics={},
        generated_at=datetime(2026, 1, 1),
        generated_by="tester",
        template_filename="t",
    )


def _docx_bytes(paragraph_text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(paragraph_text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_html_sandbox_blocks_rce():
    template = f"<pre>{RCE_PAYLOAD}</pre>".encode()
    with pytest.raises(SecurityError):
        render_html_report(template, _context())


def test_docx_sandbox_blocks_rce():
    template = _docx_bytes(RCE_PAYLOAD)
    with pytest.raises(SecurityError):
        render_docx_report(template, _context(), None)


def test_html_benign_renders_and_tojson_works():
    template = (
        b"<h1>{{ assessment.name }}</h1>"
        b"<script>var d = {{ statistics | tojson }};</script>"
    )
    out = render_html_report(template, _context())

    assert "Acme Engagement" in out
    # tojson emits raw JSON (not HTML-entity-encoded) so the <script> stays valid.
    assert "var d = {};" in out


def test_docx_benign_renders():
    template = _docx_bytes("Assessment: {{ assessment.name }}")
    result = render_docx_report(template, _context(), None)

    rendered = Document(io.BytesIO(result))
    text = "\n".join(p.text for p in rendered.paragraphs)
    assert "Assessment: Acme Engagement" in text
