"""
Markdown handling for reports (DOCX path).

Single parser (markdown-it-py), used by:
- md_to_subdoc() — converts markdown to OOXML paragraphs with proper
  Word styles (Heading 1-6, List Bullet, List Number, Code, Normal).

Also provides prepare_docx_context() to walk a context dict and convert
all markdown string fields before passing to docxtpl.

DOCX template requirements:
- Markdown fields must use {{p field }} (not {{ field | rich }}) so docxtpl
  replaces the whole paragraph with the generated XML block.
- Styles 'Heading 1'–'Heading 6', 'List Bullet', 'List Number', 'List Bullet 2',
  'List Number 2', and 'Code' should be defined in the template. Missing styles
  fall back to the document default ('Normal').

Image handling:
- DOCX: Images are embedded as inline drawings via the template's document part.
  The image bytes are added as a relationship on tpl.get_docx().part and
  referenced by rId in the generated OOXML.
- HTML: Images are handled entirely client-side in the template via base64 data URIs.
"""

import re
from typing import Any

from markdown_it import MarkdownIt

md = MarkdownIt("commonmark")

# Regex to extract file ID from API image URLs
_FILE_ID_RE = re.compile(
    r"/api/v1/assessments/[^/]+/activity/[^/]+/files/([^/]+)/download"
)


def _extract_file_id(src: str) -> str | None:
    """Extract file ID from an API image URL, or return None."""
    m = _FILE_ID_RE.match(src)
    return m.group(1) if m else None


# ---------------------------------------------------------------------------
# DOCX path: markdown -> _MarkdownSubdoc (block-level, proper Word styles)
# ---------------------------------------------------------------------------


class _MarkdownSubdoc:
    """
    Builds OOXML paragraph XML for docxtpl {{p field }} injection.

    Uses python-docx's OxmlElement API for correct namespace handling —
    no docxcompose dependency required. Style IDs are resolved from the
    template document at construction time so the inserted XML correctly
    references styles already present in the template.

    DOCX templates must use {{p field }} (not {{ field | rich }}) so that
    docxtpl replaces the whole paragraph with the generated XML.
    """

    def __init__(self, tpl: Any) -> None:
        from docx import Document
        from docx.oxml.ns import qn

        self._tpl = tpl

        # Snapshot style IDs from the template (name → w:styleId)
        docx_tpl = tpl.get_docx()
        self._style_ids: dict[str, str] = {
            style.name: style.style_id for style in docx_tpl.styles
        }

        # Fresh document for XML namespace context — no _part sharing needed
        self._doc = Document()
        body = self._doc.element.body
        for p in list(body.findall(qn("w:p"))):
            body.remove(p)

        self._pic_counter = 0

    def _style_id(self, name: str) -> str | None:
        return self._style_ids.get(name)

    def new_para(self, style_name: str | None = None, indent: int = 0) -> Any:
        """Append a new paragraph to the body and return the element.

        Args:
            style_name: Word style name to apply.
            indent: Left indent in twips (1440 twips = 1 inch).
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p = OxmlElement("w:p")
        pPr: Any = None
        if style_name:
            sid = self._style_id(style_name)
            if sid:
                pPr = OxmlElement("w:pPr")
                pStyle = OxmlElement("w:pStyle")
                pStyle.set(qn("w:val"), sid)
                pPr.append(pStyle)
        if indent > 0:
            if pPr is None:
                pPr = OxmlElement("w:pPr")
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(indent))
            pPr.append(ind)
        if pPr is not None:
            p.append(pPr)
        self._doc.element.body.append(p)
        return p

    def add_run(
        self,
        para: Any,
        text: str,
        bold: bool = False,
        italic: bool = False,
        font: str | None = None,
        size: int | None = None,
        style: str | None = None,
    ) -> None:
        """Append a text run with optional formatting to a paragraph element."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        r = OxmlElement("w:r")
        rpr_children: list = []
        if style:
            sid = self._style_id(style)
            if sid:
                rStyle = OxmlElement("w:rStyle")
                rStyle.set(qn("w:val"), sid)
                rpr_children.append(rStyle)
        if bold:
            rpr_children.append(OxmlElement("w:b"))
            rpr_children.append(OxmlElement("w:bCs"))
        if italic:
            rpr_children.append(OxmlElement("w:i"))
        if font:
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), font)
            rFonts.set(qn("w:hAnsi"), font)
            rpr_children.append(rFonts)
        if size:
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(size))
            szCs = OxmlElement("w:szCs")
            szCs.set(qn("w:val"), str(size))
            rpr_children.append(sz)
            rpr_children.append(szCs)
        if rpr_children:
            rPr = OxmlElement("w:rPr")
            for child in rpr_children:
                rPr.append(child)
            r.append(rPr)
        t = OxmlElement("w:t")
        if text != text.strip():
            t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        para.append(r)

    def add_break(self, para: Any) -> None:
        """Append a line-break run to a paragraph element."""
        from docx.oxml import OxmlElement

        r = OxmlElement("w:r")
        r.append(OxmlElement("w:br"))
        para.append(r)

    def add_image(self, para: Any, image_bytes: bytes) -> bool:
        """Embed an inline image into a paragraph. Returns True on success."""
        from io import BytesIO

        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls
        from docx.shared import Inches

        doc_part = self._tpl.get_docx().part
        rId, image = doc_part.get_or_add_image(BytesIO(image_bytes))
        cx, cy = int(image.width), int(image.height)

        # Cap width at 5.5 inches to fit within page margins
        max_cx = int(Inches(5.5))
        if cx > max_cx:
            cy = int(cy * max_cx / cx)
            cx = max_cx

        self._pic_counter += 1
        pic_id = self._pic_counter
        ns = nsdecls("w", "wp", "a", "pic", "r")

        drawing_xml = (
            f"<w:drawing {ns}>"
            f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
            f'<wp:extent cx="{cx}" cy="{cy}"/>'
            f'<wp:docPr id="{pic_id}" name="Picture {pic_id}"/>'
            f"<wp:cNvGraphicFramePr>"
            f'<a:graphicFrameLocks noChangeAspect="1"/>'
            f"</wp:cNvGraphicFramePr>"
            f"<a:graphic>"
            f"<a:graphicData"
            f' uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
            f"<pic:pic>"
            f"<pic:nvPicPr>"
            f'<pic:cNvPr id="0" name="image"/>'
            f"<pic:cNvPicPr/>"
            f"</pic:nvPicPr>"
            f"<pic:blipFill>"
            f'<a:blip r:embed="{rId}"/>'
            f"<a:stretch><a:fillRect/></a:stretch>"
            f"</pic:blipFill>"
            f"<pic:spPr>"
            f"<a:xfrm>"
            f'<a:off x="0" y="0"/>'
            f'<a:ext cx="{cx}" cy="{cy}"/>'
            f"</a:xfrm>"
            f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
            f"</pic:spPr>"
            f"</pic:pic>"
            f"</a:graphicData>"
            f"</a:graphic>"
            f"</wp:inline>"
            f"</w:drawing>"
        )
        drawing = parse_xml(drawing_xml)

        r = OxmlElement("w:r")
        r.append(drawing)
        para.append(r)
        return True

    def __str__(self) -> str:
        """Return inner body XML for docxtpl {{p field }} injection."""
        from lxml import etree

        body = self._doc.element.body
        if body.sectPr is not None:
            body.remove(body.sectPr)
        # Serialize each paragraph independently so namespace declarations
        # (wp:, a:, pic:, r:) are preserved on the elements that use them
        # rather than being lost when stripping the <w:body> wrapper.
        return "".join(etree.tostring(child, encoding="unicode") for child in body)


def _add_inline_to_para(
    para: Any,
    children: list,
    sd: _MarkdownSubdoc,
    image_data: dict[str, tuple[str, bytes]] | None = None,
) -> None:
    """Walk inline markdown token children and add runs to a paragraph element."""
    bold = False
    italic = False

    for child in children:
        t = child.type
        if t == "strong_open":
            bold = True
        elif t == "strong_close":
            bold = False
        elif t == "em_open":
            italic = True
        elif t == "em_close":
            italic = False
        elif t in ("link_open", "link_close"):
            pass  # Links not rendered in DOCX
        elif t == "text":
            sd.add_run(para, child.content, bold=bold, italic=italic)
        elif t == "code_inline":
            sd.add_run(
                para, child.content, style="Code-Inline", font="Courier New", size=18
            )
        elif t == "image":
            src = child.attrGet("src") or ""
            file_id = _extract_file_id(src)
            embedded = False
            if file_id and image_data and file_id in image_data:
                _content_type, img_bytes = image_data[file_id]
                embedded = sd.add_image(para, img_bytes)
            if not embedded:
                alt = child.content or (
                    child.children[0].content if child.children else "image"
                )
                sd.add_run(para, f"[{alt}]", italic=True)
        elif t in ("softbreak", "hardbreak"):
            sd.add_break(para)


def md_to_subdoc(
    text: str | None,
    tpl: Any,
    image_data: dict[str, tuple[str, bytes]] | None = None,
) -> Any:
    """
    Convert markdown text to a block of OOXML paragraphs with proper Word styles.

    Styles used (must exist in the DOCX template):
    - Headings:      'Heading 1' … 'Heading 6'
    - Bullet lists:  'List Bullet' (nested: 'List Bullet 2', 'List Bullet 3')
    - Ordered lists: 'List Number' (nested: 'List Number 2', 'List Number 3')
    - Code blocks:   'Code' (custom style expected in the template)
    - Normal text:   'Normal'
    - Inline:        bold, italic, inline code (Courier New 9pt)
    - Images:        embedded inline when image_data is provided, else [alt text]

    Missing styles fall back to the document default.

    IMPORTANT: Reference these fields with {{p field }} in the DOCX template
    (not {{ field | rich }}), as the value replaces the whole paragraph.
    """
    sd = _MarkdownSubdoc(tpl)

    if not text:
        return sd

    tokens = md.parse(text)
    list_stack: list[str] = []  # 'bullet' or 'ordered'
    current_style = "Normal"
    current_indent = 0  # left indent in twips for nested lists

    for token in tokens:
        t = token.type
        if t == "heading_open":
            tag = token.tag  # 'h1', 'h2', …
            level = int(tag[1]) if len(tag) == 2 and tag[1].isdigit() else 1
            current_style = f"Heading {min(level, 6)}"
            current_indent = 0
        elif t in ("heading_close", "paragraph_close"):
            current_style = "Normal"
            current_indent = 0
        elif t == "paragraph_open":
            if list_stack:
                level = len(list_stack)
                if list_stack[-1] == "bullet":
                    current_style = "TableBodyBullet"
                else:
                    current_style = "List Number"
                # Indent nested lists: 720 twips (0.5 inch) per extra level
                current_indent = (level - 1) * 720
            else:
                current_style = "Normal"
                current_indent = 0
        elif t == "bullet_list_open":
            list_stack.append("bullet")
        elif t == "bullet_list_close":
            if list_stack:
                list_stack.pop()
        elif t == "ordered_list_open":
            list_stack.append("ordered")
        elif t == "ordered_list_close":
            if list_stack:
                list_stack.pop()
        elif t in ("fence", "code_block"):
            for line in token.content.rstrip("\n").split("\n"):
                para = sd.new_para("Code")
                sd.add_run(para, line, font="Courier New", size=18)
        elif t == "inline" and token.children:
            para = sd.new_para(current_style, indent=current_indent)
            _add_inline_to_para(para, token.children, sd, image_data)

    return sd


# ---------------------------------------------------------------------------
# DOCX context preparation
# ---------------------------------------------------------------------------

# Fields in ActivityReport that contain markdown
MARKDOWN_FIELDS = {
    "rationale",
    "actions",
    "requirements",
    "notes",
    "log_notes",
    "alert_notes",
    "prevent_notes",
    "stakeholder_notification_notes",
}


def prepare_docx_context(
    context: dict,
    tpl: Any,
    image_data: dict[str, tuple[str, bytes]] | None = None,
) -> dict:
    """
    Walk a ReportContext dict (from dataclasses.asdict) and convert
    markdown string fields to Subdoc objects for DOCX rendering.

    Handles both activities_grouped and activities_flat lists.

    Args:
        context: The report context dict from asdict(ReportContext).
        tpl: The DocxTemplate instance (required for OOXML generation).
        image_data: Dict of file_id -> (content_type, bytes) for image embedding.
    """
    # Process flat activities
    if "activities_flat" in context:
        for activity in context["activities_flat"]:
            _convert_markdown_fields(activity, tpl, image_data)

    # Process grouped activities
    if "activities_grouped" in context:
        for group in context["activities_grouped"]:
            for activity in group.get("activities", []):
                _convert_markdown_fields(activity, tpl, image_data)

    return context


def _convert_markdown_fields(
    activity: dict,
    tpl: Any,
    image_data: dict[str, tuple[str, bytes]] | None = None,
) -> None:
    """
    Convert markdown fields in an activity dict to Subdoc objects.

    Subdoc produces proper Word styles; use {{p field }} in the DOCX template.
    """
    for field_name in MARKDOWN_FIELDS:
        if field_name in activity and activity[field_name]:
            activity[field_name] = md_to_subdoc(activity[field_name], tpl, image_data)
